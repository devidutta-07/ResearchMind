import streamlit as st
import time
from agents import web_extractor_agent, web_search_agent, writer_chain, critic_chain

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ResearchMind · AI Research Agent",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Shared CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Mono:wght@300;400;500&family=DM+Sans:ital,wght@0,300;0,400;0,500;1,300&display=swap');

/* ── Palette
   bg:        #08060f  (near-black with violet undertone)
   surface:   #110d1e
   accent:    #9d6fff  (bright violet)
   accent2:   #c4a3ff  (lavender)
   silver:    #b8b4c8  (cool silver text)
   done:      #7ee8c4  (teal-mint for "done" state)
────────────────────────────────────── */

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
    color: #e4e0f0;
}
.stApp {
    background: #08060f;
    background-image:
        radial-gradient(ellipse 80% 55% at 15% -5%,  rgba(120,60,255,0.18) 0%, transparent 60%),
        radial-gradient(ellipse 65% 45% at 85% 105%, rgba(180,100,255,0.10) 0%, transparent 55%),
        radial-gradient(ellipse 40% 30% at 55% 50%,  rgba(80,30,180,0.06)  0%, transparent 50%);
}
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 2rem 3rem 4rem; max-width: 1200px; }

/* ── Hero ── */
.hero { text-align: center; padding: 3.5rem 0 2.5rem; }
.hero-eyebrow {
    font-family: 'DM Mono', monospace; font-size: 0.7rem; font-weight: 500;
    letter-spacing: 0.25em; text-transform: uppercase; color: #c4a3ff;
    margin-bottom: 1rem; opacity: 0.9;
}
.hero h1 {
    font-family: 'Syne', sans-serif;
    font-size: clamp(2.8rem, 6vw, 5rem); font-weight: 800;
    line-height: 1.0; letter-spacing: -0.03em; color: #ede8ff; margin: 0 0 1rem;
}
.hero h1 span {
    background: linear-gradient(120deg, #9d6fff, #c4a3ff);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.hero-sub {
    font-size: 1.05rem; font-weight: 300; color: #7b748f;
    max-width: 520px; margin: 0 auto; line-height: 1.65;
}
.divider {
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(157,111,255,0.35), transparent);
    margin: 2rem 0;
}

/* ── Input card ── */
.input-card {
    background: rgba(157,111,255,0.04); border: 1px solid rgba(157,111,255,0.18);
    border-radius: 16px; padding: 2rem 2.5rem; margin-bottom: 2rem;
    backdrop-filter: blur(8px);
}
.stTextInput > div > div > input {
    background: rgba(157,111,255,0.06) !important;
    border: 1px solid rgba(157,111,255,0.28) !important;
    border-radius: 10px !important; color: #ede8ff !important;
    font-family: 'DM Sans', sans-serif !important; font-size: 1rem !important;
    padding: 0.75rem 1rem !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
}
.stTextInput > div > div > input:focus {
    border-color: #9d6fff !important;
    box-shadow: 0 0 0 3px rgba(157,111,255,0.15) !important;
}
.stTextInput > label {
    font-family: 'DM Mono', monospace !important; font-size: 0.72rem !important;
    letter-spacing: 0.15em !important; text-transform: uppercase !important;
    color: #c4a3ff !important; font-weight: 500 !important;
}

/* ── Buttons ── */
.stButton > button {
    background: linear-gradient(135deg, #9d6fff 0%, #6a30e8 100%) !important;
    color: #ffffff !important; font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important; font-size: 0.95rem !important;
    letter-spacing: 0.04em !important; border: none !important;
    border-radius: 10px !important; padding: 0.7rem 2.2rem !important;
    cursor: pointer !important; transition: transform 0.15s, box-shadow 0.15s !important;
    box-shadow: 0 4px 22px rgba(157,111,255,0.38) !important; width: 100%;
}
.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 30px rgba(157,111,255,0.52) !important;
}

/* ── Step cards ── */
.step-card {
    background: rgba(255,255,255,0.025); border: 1px solid rgba(255,255,255,0.06);
    border-radius: 14px; padding: 1.5rem 1.8rem; margin-bottom: 1.2rem;
    position: relative; overflow: hidden; transition: border-color 0.3s;
}
.step-card.active { border-color: rgba(157,111,255,0.45); background: rgba(157,111,255,0.05); }
.step-card.done   { border-color: rgba(126,232,196,0.30); background: rgba(126,232,196,0.03); }
.step-card::before {
    content: ''; position: absolute; left: 0; top: 0; bottom: 0; width: 3px;
    border-radius: 14px 0 0 14px; background: rgba(255,255,255,0.05); transition: background 0.3s;
}
.step-card.active::before { background: linear-gradient(180deg, #9d6fff, #c4a3ff); }
.step-card.done::before   { background: #7ee8c4; }
.step-header { display: flex; align-items: center; gap: 0.8rem; margin-bottom: 0.3rem; }
.step-num   { font-family: 'DM Mono', monospace; font-size: 0.68rem; font-weight: 500; letter-spacing: 0.15em; color: #9d6fff; opacity: 0.8; }
.step-title { font-family: 'Syne', sans-serif; font-size: 0.95rem; font-weight: 700; color: #ede8ff; }
.step-status { margin-left: auto; font-family: 'DM Mono', monospace; font-size: 0.68rem; letter-spacing: 0.1em; }
.status-waiting { color: #3d3550; }
.status-running { color: #c4a3ff; }
.status-done    { color: #7ee8c4; }
.section-heading { font-family: 'Syne', sans-serif; font-size: 1.3rem; font-weight: 700; color: #ede8ff; margin: 2rem 0 1rem; }

/* ── Results page ── */
.results-hero { padding: 2.5rem 0 1.5rem; border-bottom: 1px solid rgba(157,111,255,0.14); margin-bottom: 2rem; }
.results-hero h2 { font-family: 'Syne', sans-serif; font-size: 2rem; font-weight: 800; color: #ede8ff; margin: 0 0 0.4rem; }
.results-hero .topic-label { font-family: 'DM Mono', monospace; font-size: 0.72rem; color: #9d6fff; letter-spacing: 0.15em; text-transform: uppercase; margin-bottom: 0.5rem; }
.results-hero .topic-name  { font-size: 1rem; color: #7b748f; }

.report-panel {
    background: rgba(157,111,255,0.03); border: 1px solid rgba(157,111,255,0.22);
    border-radius: 16px; padding: 2rem 2.5rem; margin-top: 1rem;
}
.feedback-panel {
    background: rgba(126,232,196,0.025); border: 1px solid rgba(126,232,196,0.20);
    border-radius: 16px; padding: 2rem 2.5rem; margin-top: 1rem;
}
.panel-label {
    font-family: 'DM Mono', monospace; font-size: 0.7rem;
    letter-spacing: 0.2em; text-transform: uppercase; margin-bottom: 1.2rem;
    padding-bottom: 0.7rem;
}
.panel-label.orange { color: #c4a3ff; border-bottom: 1px solid rgba(157,111,255,0.18); }
.panel-label.green  { color: #7ee8c4; border-bottom: 1px solid rgba(126,232,196,0.18); }
.result-panel {
    background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.06);
    border-radius: 14px; padding: 1.8rem 2rem; margin-top: 1rem; margin-bottom: 1.5rem;
}
.result-panel-title { font-family: 'DM Mono', monospace; font-size: 0.7rem; font-weight: 500; letter-spacing: 0.2em; text-transform: uppercase; color: #9d6fff; margin-bottom: 1rem; padding-bottom: 0.7rem; border-bottom: 1px solid rgba(157,111,255,0.15); }
.result-content { font-size: 0.92rem; line-height: 1.8; color: #b8b4c8; white-space: pre-wrap; font-family: 'DM Sans', sans-serif; }

/* ── Back button (ghost style) ── */
.stButton.back-btn > button {
    background: transparent !important;
    border: 1px solid rgba(157,111,255,0.35) !important;
    color: #c4a3ff !important; width: auto !important;
    box-shadow: none !important; font-size: 0.85rem !important;
    padding: 0.5rem 1.2rem !important;
}
.stButton.back-btn > button:hover {
    background: rgba(157,111,255,0.10) !important;
    transform: none !important; box-shadow: none !important;
}

/* ── Download buttons ── */
.dl-wrapper {
    display: inline-flex; align-items: center; gap: 0.5rem;
    background: rgba(255,255,255,0.03); border: 1px solid rgba(255,255,255,0.08);
    border-radius: 10px; padding: 0.35rem 0.8rem;
}
.stDownloadButton > button {
    background: rgba(157,111,255,0.12) !important;
    border: 1px solid rgba(157,111,255,0.32) !important;
    color: #c4a3ff !important; font-family: 'DM Mono', monospace !important;
    font-size: 0.72rem !important; letter-spacing: 0.12em !important;
    text-transform: uppercase !important; font-weight: 500 !important;
    border-radius: 8px !important; padding: 0.45rem 1rem !important;
    width: auto !important; box-shadow: none !important;
    transition: background 0.15s, border-color 0.15s !important;
}
.stDownloadButton > button:hover {
    background: rgba(157,111,255,0.24) !important;
    border-color: #9d6fff !important;
    transform: none !important; box-shadow: none !important;
}

.notice { font-family: 'DM Mono', monospace; font-size: 0.72rem; color: #3a3250; text-align: center; margin-top: 3rem; letter-spacing: 0.08em; }
</style>
""", unsafe_allow_html=True)


# ── Session state init ────────────────────────────────────────────────────────
for key, default in [("results", {}), ("running", False), ("done", False), ("page", "home"), ("topic_val", "")]:
    if key not in st.session_state:
        st.session_state[key] = default


# ── Helper: step card ─────────────────────────────────────────────────────────
def step_card(num, title, state, desc=""):
    status_map = {"waiting": ("WAITING", "status-waiting"), "running": ("● RUNNING", "status-running"), "done": ("✓ DONE", "status-done")}
    label, cls = status_map.get(state, ("", ""))
    card_cls = {"running": "active", "done": "done"}.get(state, "")
    st.markdown(f"""
    <div class="step-card {card_cls}">
        <div class="step-header">
            <span class="step-num">{num}</span>
            <span class="step-title">{title}</span>
            <span class="step-status {cls}">{label}</span>
        </div>
        {"<div style='font-size:0.82rem;color:#706860;margin-top:0.3rem;'>" + desc + "</div>" if desc else ""}
    </div>""", unsafe_allow_html=True)


# ── Helper: convert markdown → simple DOCX bytes ─────────────────────────────
def md_to_docx_bytes(md_text: str) -> bytes:
    try:
        from docx import Document
        from io import BytesIO
        doc = Document()
        for line in md_text.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return None


# ── Helper: convert markdown → simple PDF bytes ──────────────────────────────
def md_to_pdf_bytes(md_text: str, topic: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
        from io import BytesIO
        import re

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2.5*cm, rightMargin=2.5*cm,
                                topMargin=2.5*cm, bottomMargin=2.5*cm)
        styles = getSampleStyleSheet()
        story = []

        for line in md_text.split("\n"):
            if line.startswith("# "):
                story.append(Paragraph(line[2:], styles["h1"]))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["h2"]))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:], styles["h3"]))
            elif line.strip():
                clean = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", line)
                clean = re.sub(r"\*(.*?)\*",   r"<i>\1</i>", clean)
                story.append(Paragraph(clean, styles["Normal"]))
            else:
                story.append(Spacer(1, 0.3*cm))

        doc.build(story)
        return buf.getvalue()
    except ImportError:
        return None


# ════════════════════════════════════════════════════════════════════════════
#  PAGE: HOME
# ════════════════════════════════════════════════════════════════════════════
if st.session_state.page == "home":

    # ── Hero ──
    st.markdown("""
    <div class="hero">
        <div class="hero-eyebrow">Multi-Agent AI System</div>
        <h1>Research<span>Mind</span></h1>
        <p class="hero-sub">
            Four specialized AI agents collaborate — searching, scraping, writing,
            and critiquing — to deliver a polished research report on any topic.
        </p>
    </div>
    <div class="divider"></div>
    """, unsafe_allow_html=True)

    col_input, col_spacer, col_pipeline = st.columns([5, 0.5, 4])

    with col_input:
        st.markdown('<div class="input-card">', unsafe_allow_html=True)
        topic = st.text_input("Research Topic", placeholder="e.g. Quantum computing breakthroughs in 2025", key="topic_input")
        run_btn = st.button("⚡  Run Research Pipeline", use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("""<div style="display:flex;gap:0.5rem;flex-wrap:wrap;margin-bottom:1.5rem;">
            <span style="font-family:'DM Mono',monospace;font-size:0.68rem;color:#3a3250;letter-spacing:0.1em;">TRY →</span>""",
            unsafe_allow_html=True)
        for ex in ["LLM agents 2025", "CRISPR gene editing", "Fusion energy progress"]:
            st.markdown(f"""<span style="background:rgba(157,111,255,0.07);border:1px solid rgba(157,111,255,0.18);
                border-radius:6px;padding:0.25rem 0.7rem;font-size:0.75rem;color:#b8b4c8;
                font-family:'DM Sans',sans-serif;">{ex}</span>""", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    with col_pipeline:
        st.markdown('<div class="section-heading">Pipeline</div>', unsafe_allow_html=True)
        r = st.session_state.results

        def s(step):
            steps = ["search", "reader", "writer", "critic"]
            if step in r: return "done"
            if st.session_state.running:
                for k in steps:
                    if k not in r:
                        return "running" if k == step else "waiting"
            return "waiting"

        step_card("01", "Search Agent",  s("search"), "Gathers recent web information")
        step_card("02", "Reader Agent",  s("reader"), "Scrapes & extracts deep content")
        step_card("03", "Writer Chain",  s("writer"), "Drafts the full research report")
        step_card("04", "Critic Chain",  s("critic"), "Reviews & scores the report")

    # ── Trigger pipeline ──
    if run_btn:
        if not topic.strip():
            st.warning("Please enter a research topic first.")
        else:
            st.session_state.topic_val = topic.strip()
            st.session_state.results = {}
            st.session_state.running = True
            st.session_state.done = False
            st.rerun()

    if st.session_state.running and not st.session_state.done:
        results = {}
        topic_val = st.session_state.topic_val

        with st.spinner("🔍  Search Agent is working…"):
            search_agent = web_search_agent()
            sr = search_agent.invoke({"messages": [("user", f"Find recent, reliable and detailed information about: {topic_val}")]})
            results["search"] = sr["messages"][-1].content
            st.session_state.results = dict(results)

        with st.spinner("📄  Reader Agent is scraping top resources…"):
            reader_agent = web_extractor_agent()
            rr = reader_agent.invoke({"messages": [("user",
                f"Based on the following search results about '{topic_val}', "
                f"pick the most relevant URL and scrape it for deeper content.\n\nSearch Results:\n{results['search'][:800]}")]})
            results["reader"] = rr["messages"][-1].content
            st.session_state.results = dict(results)

        with st.spinner("✍️  Writer is drafting the report…"):
            research_combined = f"SEARCH RESULTS:\n{results['search']}\n\nDETAILED SCRAPED CONTENT:\n{results['reader']}"
            results["writer"] = writer_chain.invoke({"topic": topic_val, "research": research_combined})
            st.session_state.results = dict(results)

        with st.spinner("🧐  Critic is reviewing the report…"):
            results["critic"] = critic_chain.invoke({"report": results["writer"]})
            st.session_state.results = dict(results)

        st.session_state.running = False
        st.session_state.done = True
        st.session_state.page = "results"   # ← navigate to results page
        st.rerun()

    st.markdown('<div class="notice">ResearchMind · Powered by LangChain multi-agent pipeline · Built with Streamlit</div>', unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════════════════
#  PAGE: RESULTS
# ════════════════════════════════════════════════════════════════════════════
elif st.session_state.page == "results":
    r = st.session_state.results
    topic_val = st.session_state.topic_val
    ts = int(time.time())

    # ── Top bar: back + download ──
    col_back, col_spacer, col_dl = st.columns([2, 5, 3])

    with col_back:
        st.markdown('<div class="back-btn">', unsafe_allow_html=True)
        if st.button("← New Research"):
            st.session_state.page = "home"
            st.session_state.results = {}
            st.session_state.done = False
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    with col_dl:
        # Small inline download cluster
        st.markdown("""
        <div style="display:flex;align-items:center;justify-content:flex-end;gap:0.5rem;padding-top:0.4rem;">
            <span style="font-family:'DM Mono',monospace;font-size:0.65rem;color:#3a3250;letter-spacing:0.12em;">DOWNLOAD</span>
        </div>""", unsafe_allow_html=True)

        dl_col1, dl_col2, dl_col3 = st.columns(3)
        report_md = r.get("writer", "")

        with dl_col1:
            st.download_button(
                label="MD",
                data=report_md,
                file_name=f"report_{ts}.md",
                mime="text/markdown",
                use_container_width=True,
            )
        with dl_col2:
            pdf_bytes = md_to_pdf_bytes(report_md, topic_val)
            if pdf_bytes:
                st.download_button(
                    label="PDF",
                    data=pdf_bytes,
                    file_name=f"report_{ts}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            else:
                st.markdown('<span style="font-family:\'DM Mono\',monospace;font-size:0.65rem;color:#444;padding:0.45rem 0.8rem;display:block;text-align:center;">PDF*</span>', unsafe_allow_html=True)

        with dl_col3:
            docx_bytes = md_to_docx_bytes(report_md)
            if docx_bytes:
                st.download_button(
                    label="DOCX",
                    data=docx_bytes,
                    file_name=f"report_{ts}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            else:
                st.markdown('<span style="font-family:\'DM Mono\',monospace;font-size:0.65rem;color:#444;padding:0.45rem 0.8rem;display:block;text-align:center;">DOCX*</span>', unsafe_allow_html=True)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # ── Results hero ──
    st.markdown(f"""
    <div class="results-hero">
        <div class="topic-label">Research Report</div>
        <h2>Results</h2>
        <div class="topic-name">Topic: {topic_val}</div>
    </div>""", unsafe_allow_html=True)

    # ── Raw outputs (collapsed) ──
    if "search" in r:
        with st.expander("🔍 Search Results (raw)", expanded=False):
            st.markdown(f'<div class="result-panel"><div class="result-panel-title">Search Agent Output</div>'
                        f'<div class="result-content">{r["search"]}</div></div>', unsafe_allow_html=True)

    if "reader" in r:
        with st.expander("📄 Scraped Content (raw)", expanded=False):
            st.markdown(f'<div class="result-panel"><div class="result-panel-title">Reader Agent Output</div>'
                        f'<div class="result-content">{r["reader"]}</div></div>', unsafe_allow_html=True)

    # ── Final report ──
    if "writer" in r:
        st.markdown('<div class="report-panel"><div class="panel-label orange">📝 Final Research Report</div>', unsafe_allow_html=True)
        st.markdown(r["writer"])
        st.markdown("</div>", unsafe_allow_html=True)

    # ── Critic feedback ──
    if "critic" in r:
        st.markdown('<div class="feedback-panel"><div class="panel-label green">🧐 Critic Feedback</div>', unsafe_allow_html=True)
        st.markdown(r["critic"])
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="notice">ResearchMind · Powered by LangChain multi-agent pipeline · Built with Streamlit</div>', unsafe_allow_html=True)

    # ── Install note if PDF/DOCX libraries missing ──
    if md_to_pdf_bytes("test", "test") is None or md_to_docx_bytes("test") is None:
        st.markdown("""
        <div style="margin-top:1rem;font-family:'DM Mono',monospace;font-size:0.68rem;color:#3a3250;text-align:center;">
            * Install <code>reportlab</code> and <code>python-docx</code> to enable PDF & DOCX downloads
        </div>""", unsafe_allow_html=True)